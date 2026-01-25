"""
DOCX document ingestion using python-docx.
Leverages Word's built-in structure (headings, paragraphs).
"""

import re
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    Document = None

from .base import DocumentIngester, ExtractedDocument


class DOCXIngester(DocumentIngester):
    """Ingester for DOCX documents."""

    SUPPORTED_EXTENSIONS = [".docx"]

    def __init__(self, normalize_whitespace: bool = True):
        super().__init__(normalize_whitespace)

        if Document is None:
            raise ImportError(
                "python-docx is required for DOCX ingestion. Install with: pip install python-docx"
            )

    def extract(self, path: Path) -> ExtractedDocument:
        """Extract text and structure from a DOCX document."""
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {path}")

        warnings = []
        doc = Document(path)

        # Extract metadata
        core_props = doc.core_properties
        title = core_props.title or self._clean_extracted_title(path.stem)
        author = core_props.author

        # Extract text and structure
        paragraphs_text = []
        chapters = []
        current_chapter_start = 0
        current_chapter_title = None
        has_images = False

        for para in doc.paragraphs:
            para_text = para.text.strip()

            if not para_text:
                paragraphs_text.append("")
                continue

            # Check if this is a heading (potential chapter marker)
            is_heading = False
            heading_level = None

            if para.style and para.style.name:
                style_name = para.style.name.lower()
                if "heading" in style_name:
                    is_heading = True
                    # Extract heading level if present
                    level_match = re.search(r"heading\s*(\d+)", style_name)
                    if level_match:
                        heading_level = int(level_match.group(1))

            # Track chapter structure
            if is_heading and heading_level in (1, 2):
                # Close previous chapter if exists
                if current_chapter_title is not None:
                    chapters.append(
                        {
                            "title": current_chapter_title,
                            "start_pos": current_chapter_start,
                            "end_pos": len("\n\n".join(paragraphs_text)),
                            "heading_level": heading_level,
                        }
                    )

                current_chapter_title = para_text
                current_chapter_start = len("\n\n".join(paragraphs_text))

            paragraphs_text.append(para_text)

        # Close final chapter
        if current_chapter_title is not None:
            chapters.append(
                {
                    "title": current_chapter_title,
                    "start_pos": current_chapter_start,
                    "end_pos": len("\n\n".join(paragraphs_text)),
                }
            )

        # Check for images
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                has_images = True
                break

        if has_images:
            warnings.append("Document contains images that won't be included in text extraction")

        # Join paragraphs
        full_text = "\n\n".join(paragraphs_text)

        # Normalize if requested
        if self.normalize_whitespace:
            full_text = self._normalize_text(full_text)

        # If no chapters detected via headings, try regex detection
        if not chapters:
            chapters = self._detect_chapters_fallback(full_text)

        return ExtractedDocument(
            text=full_text,
            source_path=path,
            source_format="docx",
            title=title,
            author=author,
            chapters=chapters if chapters else None,
            page_count=None,  # DOCX doesn't have fixed pages
            has_images=has_images,
            extraction_warnings=warnings,
        )

    def _detect_chapters_fallback(self, text: str) -> Optional[list[dict]]:
        """Fallback chapter detection using regex patterns."""
        chapters = []

        patterns = [
            r"^(?P<title>Chapter\s+(?:\d+|[IVXLC]+|One|Two|Three|Four|Five|Six|Seven|Eight|Nine|Ten)[^\n]*)",
            r"^(?P<title>CHAPTER\s+(?:\d+|[IVXLC]+)[^\n]*)",
        ]

        for pattern in patterns:
            for match in re.finditer(pattern, text, re.MULTILINE | re.IGNORECASE):
                chapters.append(
                    {
                        "title": match.group("title").strip(),
                        "start_pos": match.start(),
                        "end_pos": None,
                    }
                )

        if not chapters:
            return None

        # Sort and fill end positions
        chapters.sort(key=lambda c: c["start_pos"])
        for i in range(len(chapters) - 1):
            chapters[i]["end_pos"] = chapters[i + 1]["start_pos"]
        chapters[-1]["end_pos"] = len(text)

        return chapters
